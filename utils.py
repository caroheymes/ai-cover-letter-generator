import os
import pathlib
import pymupdf
import pymupdf4llm
import streamlit as st
from crewai import Agent
from crewai_tools import WebsiteSearchTool, SerperDevTool, FileReadTool
#export en .docx
from io import BytesIO
from docx import Document
from docx.shared import Pt
import re




def convert_cv_to_md(uploaded_file, output_path="cv_md.md"):
    """Convertit un CV uploadé (PDF/DOCX/MD) en markdown."""
    try:
        file_extension = pathlib.Path(uploaded_file.name).suffix.lower()
        
        if file_extension == '.pdf':
            # Sauvegarder temporairement
            temp_path = f"{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            # Convertir
            md_text = pymupdf4llm.to_markdown(temp_path)
            pathlib.Path(output_path).write_bytes(md_text.encode("utf-8"))
            
            # Nettoyer
            os.remove(temp_path)
            
        elif file_extension == '.docx':
            temp_path = f"{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            # Convertir via PyMuPDF
            doc = pymupdf.open(temp_path)
            pdf_bytes = doc.convert_to_pdf()
            pdf_doc = pymupdf.open("pdf", pdf_bytes)
            md_text = pymupdf4llm.to_markdown(pdf_doc)
            pathlib.Path(output_path).write_bytes(md_text.encode("utf-8"))
            
            # Nettoyer
            os.remove(temp_path)
            
        elif file_extension == '.md':
            # Directement sauvegarder
            pathlib.Path(output_path).write_bytes(uploaded_file.getvalue())
            
        else:
            raise ValueError(f"Format non supporté: {file_extension}")
            
        return output_path
        
    except Exception as e:
        st.error(f"Erreur de conversion: {str(e)}")
        return None

def load_agents_from_yaml(yaml_path, temperatures, context=None):
    """Charge les agents depuis YAML avec températures dynamiques."""
    import yaml
    from crewai import Agent
    
    with open(yaml_path, 'r', encoding='utf-8') as f:
        agents_config = yaml.safe_load(f)

    context = context or {}
    
    # Map tool names to actual tool instances
    tools_map = {
        "WebsiteSearchTool": WebsiteSearchTool(),
        "SerperDevTool": SerperDevTool(),
        "FileReadTool": FileReadTool(file_path="cv_md.md")
    }
    
    agents = {}
    for agent_name, config in agents_config.items():

        # Formater les strings avec le contexte (candidate_profile)
        for key in ['role', 'goal', 'backstory']:
            if key in config and isinstance(config[key], str):
                config[key] = config[key].format(**context)
        
        # Map tool names to actual tool instances
        agent_tools = [tools_map[tool] for tool in config.get("tools", []) if tool in tools_map]
        
        agents[agent_name] = Agent(
            role=config["role"],
            goal=config["goal"],
            backstory=config["backstory"],
            tools=agent_tools,  # ✅ Pass the actual tool instances
            temperature=temperatures.get(agent_name, 0.5),
            verbose=config.get("verbose", True),
            max_iter=config.get("max_iter", 3),
            allow_delegation=config.get("allow_delegation", False)
        )
    
    return agents

def load_tasks_from_yaml(filepath, context=None):
    import yaml
    from crewai import Task
    from textwrap import dedent
    
    context = context or {}
    
    with open(filepath, 'r') as f:
        tasks_config = yaml.safe_load(f)
    
    tasks = []
    for name, config in tasks_config.items():
        # Formater les strings avec le contexte
        for key in ['description', 'expected_output']:
            if key in config and isinstance(config[key], str):
                config[key] = config[key].format(**context)
        
        task = Task(
            description=config['description'],
            expected_output=config['expected_output'],
            async_execution=config.get('async_execution', False),
            agent=None  # Sera assigné plus tard
        )
        tasks.append(task)
    
    return tasks

# def create_docx(markdown_text):
#     """
#     Convertit le texte Markdown en fichier Word (.docx) en mémoire.
#     Nettoie la syntaxe Markdown de base (**gras**, # Titres) pour un rendu propre.
#     """
#     doc = Document()
    
#     # Définir le style de base (Police standard)
#     style = doc.styles['Normal']
#     font = style.font
#     font.name = 'Calibri'
#     font.size = Pt(11)

#     # Séparer le texte par paragraphes (saut de ligne double ou simple)
#     # On normalise les sauts de ligne
#     paragraphs = markdown_text.split('\n')
    
#     for para in paragraphs:
#         # Ignorer les lignes vides
#         if not para.strip():
#             continue
            
#         # Nettoyage basique du Markdown
#         clean_text = para.strip()
        
#         # Gestion des Titres (## Titre) -> On les met en gras
#         is_heading = False
#         if clean_text.startswith('#'):
#             clean_text = clean_text.lstrip('#').strip()
#             is_heading = True
            
#         # Suppression du gras Markdown (**texte**) pour le mettre proprement
#         # Note: Pour faire simple ici, on retire juste les étoiles. 
#         # Une gestion parfaite du gras demanderait un parser plus complexe.
#         clean_text = clean_text.replace('**', '').replace('__', '')
        
#         # Ajouter le paragraphe au document
#         p = doc.add_paragraph(clean_text)
        
#         # Appliquer un formatage si c'était un titre ou un champ important
#         if is_heading:
#             p.runs[0].bold = True
#             p.runs[0].font.size = Pt(12)

#     # Sauvegarder dans un buffer mémoire (pas de fichier sur le disque)
#     bio = BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

def _add_markdown_to_doc(doc, text):
    """Fonction interne pour nettoyer et ajouter du markdown dans un objet docx"""
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        clean_text = para.strip()
        is_heading = False
        
        # Gestion Titres
        if clean_text.startswith('#'):
            clean_text = clean_text.lstrip('#').strip()
            is_heading = True
            
        # Nettoyage gras basique
        clean_text = clean_text.replace('**', '').replace('__', '')
        
        p = doc.add_paragraph(clean_text)
        
        if is_heading:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
            # Ajout d'un petit espacement après les titres
            p.paragraph_format.space_after = Pt(6)

def create_docx(final_text, draft_text=None, params=None):
    """Génère le DOCX avec la version finale et optionnellement le brouillon."""
    doc = Document()
    
    # Style global
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Ajouter la version FINALE
    doc.add_heading('Lettre de Motivation (Version Finale)', 0)
    
    doc.add_heading('Paramètres de la session', 0)
    doc.add_paragraph("Voici les configurations utilisées pour générer cette lettre :")
    
    # Créer un tableau
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid' # Ajoute des bordures visibles
    
    # En-tête du tableau
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Paramètre'
    hdr_cells[1].text = 'Valeur'
    
    # Remplir le tableau
    for key, value in params.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        # On convertit en string et on limite la longueur si c'est trop long (optionnel)
        row_cells[1].text = str(value)

    _add_markdown_to_doc(doc, final_text)

    # 2. Ajouter le BROUILLON (si demandé)
    if draft_text:
        # Saut de page
        doc.add_page_break()   
        
        doc.add_heading('Premier Jet (Version non-relue)', 0)
        p = doc.add_paragraph("Ceci est la version brute générée par le rédacteur avant le passage du relecteur.")
        p.italic = True
        
        _add_markdown_to_doc(doc, draft_text)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio



def show_buy_me_coffee():
    username = "CaroMS"
    button_html = f"""
    <div style="text-align: center; margin-top: 20px; margin-bottom: 20px;">
        <a href="https://www.buymeacoffee.com/{username}" target="_blank">
            <img src="https://cdn.buymeacoffee.com/buttons/v2/default-yellow.png" 
                 alt="Buy Me A Coffee" 
                 style="height: 35px !important; width: auto !important;" >
        </a>
    </div>
    """
    st.markdown(button_html, unsafe_allow_html=True)